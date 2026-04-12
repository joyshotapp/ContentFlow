"""
ContentFlow 完整 Pipeline 腳本
==============================
模擬 UI 操作：選題研究 → AI 寫文 → 事實查核 → 輸出 Markdown

Usage:
    python scripts/run_article_pipeline.py --seqno 4
"""

import asyncio
import json
import os
import re
import sqlite3
import sys
import time
from datetime import datetime
from pathlib import Path

# ── 設定 Python path ──────────────────────────────────────────
ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT / "src"))

from contentflow.agents.factcheck_agent import run_factcheck_agent
from contentflow.agents.seo_check_agent import run_seo_check_agent, suggest_internal_links
from contentflow.agents.research_agent import ResearchReport, run_research_agent
from contentflow.agents.seo_qa_agent import run_seo_qa_agent
from contentflow.agents.strategy_agent import run_strategy_agent
from contentflow.agents.writing_agent import run_writing_agent
from contentflow.config import settings
from contentflow.project_context import load_project_context, project_uses_pubmed

DB_PATH = ROOT / "data" / "contentflow.db"
OUTPUT_DIR = ROOT / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

# ── ANSI 顏色 ─────────────────────────────────────────────────
GREEN  = "\033[92m"
YELLOW = "\033[93m"
BLUE   = "\033[94m"
CYAN   = "\033[96m"
RED    = "\033[91m"
BOLD   = "\033[1m"
RESET  = "\033[0m"

def banner(step: int, title: str):
    labels = ["", "① 選題研究", "② AI 寫文", "③ 事實查核", "④ 取得成果"]
    colors = ["", BLUE, YELLOW, GREEN, CYAN]
    c = colors[step]
    print(f"\n{c}{BOLD}{'='*60}{RESET}")
    print(f"{c}{BOLD}  STEP {step}/4  {labels[step]}  —  {title}{RESET}")
    print(f"{c}{BOLD}{'='*60}{RESET}\n")

def tick(msg: str):
    print(f"  {GREEN}✓{RESET}  {msg}")

def info(msg: str):
    print(f"  {CYAN}·{RESET}  {msg}")

def warn(msg: str):
    print(f"  {YELLOW}!{RESET}  {msg}")


# ─────────────────────────────────────────────────────────────
def _clean_keyword_text(text: str) -> str:
    """清理單一關鍵字中的搜量與多餘空白。"""
    cleaned = str(text or "")
    cleaned = re.sub(r'\s*[\(（]\d+[\)）]\s*', ' ', cleaned)
    cleaned = re.sub(r'\s+\d+$', '', cleaned)
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def _split_keywords(text: str) -> list[str]:
    """將多行或逗號分隔的關鍵字欄位拆成乾淨清單。"""
    tokens = re.split(r'[\n,，]+', str(text or ''))
    items = []
    seen = set()
    for token in tokens:
        cleaned = _clean_keyword_text(token)
        if not cleaned or cleaned in seen:
            continue
        seen.add(cleaned)
        items.append(cleaned)
    return items


def _normalize_article_keywords(primary_raw: str, secondary_raw: str) -> tuple[str, list[str]]:
    """兼容舊資料：若主關鍵字欄位誤存多行，第一行視為主關鍵字，其餘併入副關鍵字。"""
    primary_parts = _split_keywords(primary_raw)
    secondary_parts = _split_keywords(secondary_raw)

    primary_keyword = primary_parts[0] if primary_parts else ""
    secondary_keywords = []
    seen = {primary_keyword} if primary_keyword else set()

    for keyword in primary_parts[1:] + secondary_parts:
        if keyword and keyword not in seen:
            seen.add(keyword)
            secondary_keywords.append(keyword)

    return primary_keyword, secondary_keywords


def load_article(seqno: int, project_id: int | None = None) -> dict:
    """從 SQLite 載入文章資料 + 策略欄位（來自 ContentCalendar）"""
    conn = sqlite3.connect(DB_PATH)
    if project_id:
        row = conn.execute(
            "SELECT id, seqno, primary_keyword, secondary_keywords, title, outline, status, article_type "
            "FROM articles WHERE seqno=? AND project_id=?", (seqno, project_id)
        ).fetchone()
    else:
        row = conn.execute(
            "SELECT id, seqno, primary_keyword, secondary_keywords, title, outline, status, article_type "
            "FROM articles WHERE seqno=?", (seqno,)
        ).fetchone()
    if not row:
        conn.close()
        raise ValueError(f"找不到 seqno={seqno} 的文章")

    article_id = row[0]
    primary_keyword, secondary_keywords = _normalize_article_keywords(row[2], row[3])

    # 嘗試從 content_calendar 載入 SEO 專員的策略指引
    strategy_context = {}
    if project_id:
        cal_row = conn.execute(
            "SELECT search_intent, target_audience, writing_architecture, faq_questions "
            "FROM content_calendar WHERE article_id=? AND project_id=?",
            (article_id, project_id),
        ).fetchone()
    else:
        cal_row = conn.execute(
            "SELECT search_intent, target_audience, writing_architecture, faq_questions "
            "FROM content_calendar WHERE article_id=?", (article_id,)
        ).fetchone()
    if not cal_row:
        # fallback: 在 calendar 的 title 或 keywords 欄位搜尋主關鍵字
        kw = primary_keyword
        if project_id:
            cal_row = conn.execute(
                "SELECT search_intent, target_audience, writing_architecture, faq_questions "
                "FROM content_calendar WHERE project_id=? AND (title LIKE ? OR keywords LIKE ?) LIMIT 1",
                (project_id, f"%{kw}%", f"%{kw}%"),
            ).fetchone()
        else:
            cal_row = conn.execute(
                "SELECT search_intent, target_audience, writing_architecture, faq_questions "
                "FROM content_calendar WHERE title LIKE ? OR keywords LIKE ? LIMIT 1",
                (f"%{kw}%", f"%{kw}%")
            ).fetchone()
    if not cal_row and len(primary_keyword) >= 3:
        # fallback 2: 用關鍵字的核心詞（去頭尾）再搜一次
        core = primary_keyword[1:-1] if len(primary_keyword) > 3 else primary_keyword
        if project_id:
            cal_row = conn.execute(
                "SELECT search_intent, target_audience, writing_architecture, faq_questions "
                "FROM content_calendar WHERE project_id=? AND (title LIKE ? OR keywords LIKE ?) LIMIT 1",
                (project_id, f"%{core}%", f"%{core}%"),
            ).fetchone()
        else:
            cal_row = conn.execute(
                "SELECT search_intent, target_audience, writing_architecture, faq_questions "
                "FROM content_calendar WHERE title LIKE ? OR keywords LIKE ? LIMIT 1",
                (f"%{core}%", f"%{core}%")
            ).fetchone()
    if not cal_row and len(primary_keyword) >= 4:
        # fallback 3: 嘗試提取 2 字核心詞（去除常見前後綴如「長/的/怎麼辦」）
        import re as _re
        core2 = _re.sub(r'^(長|膝蓋|右|左)', '', primary_keyword)
        core2 = _re.sub(r'(怎麼辦|原因|症狀|可以|不能|會好嗎)$', '', core2)
        if core2 and core2 != primary_keyword:
            if project_id:
                cal_row = conn.execute(
                    "SELECT search_intent, target_audience, writing_architecture, faq_questions "
                    "FROM content_calendar WHERE project_id=? AND (title LIKE ? OR keywords LIKE ?) LIMIT 1",
                    (project_id, f"%{core2}%", f"%{core2}%"),
                ).fetchone()
            else:
                cal_row = conn.execute(
                    "SELECT search_intent, target_audience, writing_architecture, faq_questions "
                    "FROM content_calendar WHERE title LIKE ? OR keywords LIKE ? LIMIT 1",
                    (f"%{core2}%", f"%{core2}%")
                ).fetchone()
    if cal_row:
        strategy_context = {
            "search_intent": cal_row[0] or "",
            "target_audience": cal_row[1] or "",
            "writing_architecture": cal_row[2] or "",
            "faq_questions": cal_row[3] or "",
        }
        # 過濾掉全空值
        strategy_context = {k: v for k, v in strategy_context.items() if v}

    conn.close()
    return {
        "id": article_id, "seqno": row[1], "primary_keyword": primary_keyword,
        "secondary_keywords": "\n".join(secondary_keywords), "title": row[4],
        "outline": row[5], "status": row[6],
        "article_type": row[7] or "educational",
        "strategy_context": strategy_context or None,
    }


def save_to_db(article_id: int, report_json: str, draft_content: str,
               draft_title: str, status: str,
               primary_keyword: str | None = None,
               secondary_keywords: str | None = None,
               slug: str = "",
               meta_title: str = "",
               meta_description: str = "",
               faq_schema_json: str = "",
               article_schema_json: str = "",
               seo_score: int | None = None):
    """將三個 Agent 的產出存回資料庫（與 UI 操作完全一致）"""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        UPDATE articles SET
            research_report_json = ?,
            draft_content = ?,
            title = COALESCE(NULLIF(?, ''), title),
            primary_keyword = COALESCE(NULLIF(?, ''), primary_keyword),
            secondary_keywords = COALESCE(?, secondary_keywords),
            status = ?,
            slug = ?,
            meta_title = ?,
            meta_description = ?,
            faq_schema_json = ?,
            article_schema_json = ?,
            seo_score = ?
        WHERE id = ?
    """, (
        report_json,
        draft_content,
        draft_title,
        primary_keyword,
        secondary_keywords,
        status,
        slug,
        meta_title,
        meta_description,
        faq_schema_json,
        article_schema_json,
        seo_score,
        article_id,
    ))
    conn.commit()
    conn.close()


def clean_code_fences(text: str) -> str:
    """去除 GPT 可能加上的 code fence"""
    text = re.sub(r'```markdown\n?', '\n\n', text)
    text = re.sub(r'```\n?', '\n\n', text)
    return re.sub(r'\n{3,}', '\n\n', text).strip()


# ─────────────────────────────────────────────────────────────
def export_docx(md_content: str, draft, items: list, out_path: Path) -> None:
    """將 Markdown 輸出為 Word .docx（含事實查核附錄）"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── 文章標題 ──────────────────────────────────────────────
    doc.add_heading(draft.title, level=0)

    # ── Meta 資訊 ──────────────────────────────────────────────
    for label, val in [
        ("Meta Title", draft.meta_title),
        ("Meta Description", draft.meta_description),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}：").bold = True
        p.add_run(val)

    # ── 查核摘要 ──────────────────────────────────────────────
    passed_c  = len([i for i in items if not i.needs_review])
    flagged_c = len([i for i in items if i.needs_review])
    p = doc.add_paragraph()
    p.add_run("AI 事實查核：").bold = True
    p.add_run(f"{passed_c} 項通過 ✅  /  {flagged_c} 項需人工審核 ⚠️")

    doc.add_paragraph("─" * 50)

    # ── 解析 Markdown 主文 ────────────────────────────────────
    def add_inline_md(para, text: str):
        """處理行內 **bold** 語法"""
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for j, part in enumerate(parts):
            run = para.add_run(part)
            if j % 2 == 1:
                run.bold = True

    for line in md_content.split('\n'):
        line = line.rstrip()
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline_md(p, line[2:])
        elif line.strip():
            p = doc.add_paragraph()
            add_inline_md(p, line)

    # ── 事實查核附錄（分頁）───────────────────────────────────
    doc.add_page_break()
    doc.add_heading('附錄：AI 事實查核詳情', level=1)
    doc.add_paragraph(
        f"共查核 {len(items)} 項內容宣稱，"
        f"{passed_c} 項通過，{flagged_c} 項需人工審核。"
    )

    for idx, item in enumerate(items, 1):
        icon = "✅" if not item.needs_review else "⚠️"
        p = doc.add_paragraph()
        p.add_run(f"{icon} [{idx}] ").bold = True
        p.add_run(item.claim or "")

        if item.reviewer_note:
            rp = doc.add_paragraph(f"    建議：{item.reviewer_note}")
            rp.paragraph_format.left_indent = Inches(0.3)
        if item.supporting_evidence:
            ep = doc.add_paragraph(f"    文獻：{', '.join(item.supporting_evidence)}")
            ep.paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph("\n本文章由 ContentFlow AI 自動生成，建議人工最終審閱後發佈。")
    doc.save(str(out_path))


def export_pdf(md_content: str, draft, items: list, out_path: Path) -> None:
    """將 Markdown 輸出為 PDF（透過 HTML → weasyprint）"""
    import os as _os
    # macOS homebrew 的 glib/pango 函式庫路徑
    if sys.platform == "darwin":
        brew_lib = "/opt/homebrew/lib"
        cur = _os.environ.get("DYLD_LIBRARY_PATH", "")
        if brew_lib not in cur:
            _os.environ["DYLD_LIBRARY_PATH"] = f"{brew_lib}:{cur}" if cur else brew_lib
    try:
        import markdown as md_lib
        from weasyprint import HTML
    except ImportError:
        raise ImportError("缺少套件，請執行：pip install markdown weasyprint  且  brew install pango")

    passed_c  = len([i for i in items if not i.needs_review])
    flagged_c = len([i for i in items if i.needs_review])

    # ── 主文 Markdown → HTML ──────────────────────────────────
    body_html = md_lib.markdown(
        md_content,
        extensions=['extra', 'tables', 'nl2br'],
    )

    # ── 事實查核附錄表格 ──────────────────────────────────────
    fc_rows = ""
    for idx, item in enumerate(items, 1):
        icon  = "✅" if not item.needs_review else "⚠️"
        ev    = ', '.join(item.supporting_evidence) if item.supporting_evidence else "—"
        note  = item.reviewer_note or "—"
        row_cls = "pass" if not item.needs_review else "flag"
        fc_rows += (
            f'<tr class="{row_cls}">'
            f'<td>{idx}</td><td>{icon}</td>'
            f'<td>{item.claim}</td>'
            f'<td>{note}</td>'
            f'<td>{ev}</td></tr>\n'
        )

    html = f"""<!DOCTYPE html>
<html lang="zh-TW"><head><meta charset="utf-8">
<style>
  body {{
    font-family: "PingFang TC", "Heiti TC", "Microsoft JhengHei", Arial, sans-serif;
    font-size: 13px; line-height: 1.9; margin: 0; color: #333;
  }}
  @page {{ margin: 2cm 2.5cm; size: A4; }}
  h1 {{ font-size: 22px; border-bottom: 2px solid #2563eb;
        padding-bottom: 6px; color: #1e3a8a; margin-top: 0; }}
  h2 {{ font-size: 17px; color: #2563eb; margin-top: 22px; }}
  h3 {{ font-size: 14px; color: #374151; }}
  .meta {{
    background: #eff6ff; border-left: 4px solid #2563eb;
    padding: 12px 16px; margin: 16px 0; font-size: 12px; line-height: 1.8;
  }}
  .fc-summary {{
    background: #f0fdf4; border-left: 4px solid #16a34a;
    padding: 10px 16px; margin: 16px 0; font-weight: bold; font-size: 13px;
  }}
  ul {{ margin: 6px 0; padding-left: 22px; }}
  li {{ margin-bottom: 4px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 16px 0; font-size: 11px; }}
  th {{ background: #2563eb; color: white; padding: 8px; text-align: left; }}
  td {{ padding: 6px 8px; border: 1px solid #e5e7eb; vertical-align: top; }}
  tr.pass td {{ background: #f0fdf4; }}
  tr.flag td {{ background: #fffbeb; }}
  hr {{ border: none; border-top: 1px solid #e5e7eb; margin: 28px 0; }}
  .footer {{
    color: #9ca3af; font-size: 10px; text-align: center; margin-top: 24px;
  }}
  strong {{ font-weight: bold; }}
</style></head>
<body>
<h1>{draft.title}</h1>

<div class="meta">
  <strong>Meta Title：</strong>{draft.meta_title}<br>
  <strong>Meta Description：</strong>{draft.meta_description}
</div>

<div class="fc-summary">
  AI 事實查核：{passed_c} 項通過 ✅ &nbsp;/&nbsp; {flagged_c} 項需人工審核 ⚠️
</div>

{body_html}

<hr>
<h2>附錄：AI 事實查核詳情</h2>
<table>
  <tr><th>#</th><th>狀態</th><th>宣稱內容</th><th>審核建議</th><th>支持文獻</th></tr>
  {fc_rows}
</table>

<div class="footer">本文章由 ContentFlow AI 自動生成，建議人工最終審閱後發佈。</div>
</body></html>"""

    HTML(string=html).write_pdf(str(out_path))


# ─────────────────────────────────────────────────────────────
async def run_pipeline(seqno: int, project_slug: str = "hanben"):
    t0 = time.time()

    print(f"\n{BOLD}{CYAN}ContentFlow AI Pipeline{RESET}")
    print(f"{CYAN}{'─'*40}{RESET}")

    # ── 載入專案 ──────────────────────────────────────────────
    ctx = load_project_context(project_slug=project_slug)
    print(f"  專案：{BOLD}{ctx.name}{RESET} ({ctx.slug})")
    if ctx.brand_name:
        print(f"  品牌：{ctx.brand_name}")
    if ctx.industry:
        print(f"  產業：{ctx.industry}")

    # ── 載入文章 ──────────────────────────────────────────────
    article = load_article(seqno, project_id=ctx.project_id)
    title = article["title"] or article["primary_keyword"]
    primary_kw = article["primary_keyword"]
    secondary_raw = article["secondary_keywords"] or ""

    secondary_kwds = _split_keywords(secondary_raw)

    print(f"\n  文章序號：{BOLD}#{seqno}{RESET}")
    print(f"  主關鍵字：{BOLD}{primary_kw}{RESET}")
    print(f"  副關鍵字：{secondary_kwds}")
    print(f"  標題：{title}")
    print(f"  現有狀態：{article['status']}")

    strategy = article.get("strategy_context")
    if strategy:
        print(f"\n  {CYAN}📋 SEO 策略指引（來自 SEO 專員）{RESET}")
        if strategy.get("search_intent"):
            print(f"  搜尋意圖：{strategy['search_intent']}")
        if strategy.get("target_audience"):
            print(f"  讀者痛點：{strategy['target_audience'][:60]}")
        if strategy.get("writing_architecture"):
            print(f"  架構策略：{strategy['writing_architecture'][:60]}")
        if strategy.get("faq_questions"):
            print(f"  建議 FAQ：{strategy['faq_questions'][:60]}…")
    else:
        print(f"\n  {YELLOW}（無人工策略 — STEP 1.5 將自動產生 AI 策略）{RESET}")

    # ══════════════════════════════════════════════════════════
    # STEP 1 — 選題研究
    # ══════════════════════════════════════════════════════════
    banner(1, primary_kw)

    search_keywords = secondary_kwds or [primary_kw]

    info(f"搜尋關鍵字：{search_keywords}")
    info("呼叫 Google SERP…（約 30~60 秒）")

    t1 = time.time()
    use_pubmed = project_uses_pubmed(ctx)
    report: ResearchReport = await run_research_agent(
        article_title=title,
        search_keywords=search_keywords,
        serp_gl=ctx.serp_gl,
        serp_hl=ctx.serp_hl,
        use_pubmed=use_pubmed,
    )
    elapsed1 = time.time() - t1

    pubmed_count = sum(len(r.articles) for r in report.pubmed_results)
    tick(f"研究完成！耗時 {elapsed1:.1f}s")
    tick(f"PubMed 文獻：{pubmed_count} 篇")
    tick(f"建議關鍵字：{len(report.suggested_keywords)} 個 → {', '.join(report.suggested_keywords[:6])}")
    tick(f"競品分析：{len(report.serp_analysis.top_results) if report.serp_analysis else 0} 筆")
    if report.serp_analysis and report.serp_analysis.top_results:
        for i, r in enumerate(report.serp_analysis.top_results[:3], 1):
            info(f"  競品#{i}: {r.title[:50]}")

    # ══════════════════════════════════════════════════════════
    # STEP 1.5 — 策略分析（當無人工策略時自動產生）
    # ══════════════════════════════════════════════════════════
    strategy = article.get("strategy_context")
    if not strategy:
        print(f"\n{YELLOW}{BOLD}{'─'*60}{RESET}")
        print(f"{YELLOW}{BOLD}  STEP 1.5  ⚙ AI 策略分析  —  {primary_kw}{RESET}")
        print(f"{YELLOW}{BOLD}{'─'*60}{RESET}\n")

        info("無人工策略指引，啟動 AI Strategy Agent 自動分析…")
        info(f"利用研究步驟的 SERP + PAA 資料，呼叫 {settings.llm_lite_model}…")

        t_strat = time.time()
        serp_data = report.serp_analysis
        paa_qs = report.paa_questions or []
        strategy_report = await run_strategy_agent(
            keyword=primary_kw,
            secondary_keywords=secondary_kwds,
            serp=serp_data,
            paa_questions=paa_qs,
            project_id=ctx.project_id,
        )
        elapsed_strat = time.time() - t_strat

        strategy = strategy_report.to_strategy_context()
        tick(f"策略分析完成！耗時 {elapsed_strat:.1f}s  信心度 {strategy_report.confidence}")
        info(f"搜尋意圖：{strategy_report.search_intent}")
        info(f"目標讀者：{strategy_report.target_audience[:60]}")
        info(f"架構建議：{strategy_report.writing_architecture[:60]}")
        info(f"FAQ 建議：{len(strategy_report.faq_questions)} 題")
        if strategy_report.competitor_gap:
            info(f"競品差異：{strategy_report.competitor_gap[:60]}")
    else:
        print(f"\n  {CYAN}📋 使用 SEO 專員人工策略（優先於 AI 自動分析）{RESET}")

    # ══════════════════════════════════════════════════════════
    # STEP 2 — AI 寫文
    # ══════════════════════════════════════════════════════════
    banner(2, f"依研究成果生成文章")

    writing_arch = article["outline"] or ""
    if strategy and strategy.get("writing_architecture"):
        # 策略指引中的架構優先於 outline 欄位
        writing_arch = writing_arch or strategy["writing_architecture"]
    info(f"文章架構字數：{len(writing_arch)} 字")
    if strategy:
        info("已注入 SEO 策略指引（搜尋意圖 + 讀者痛點 + FAQ）")
    info(f"呼叫 {settings.llm_lite_model} 逐段撰寫…（約 60~90 秒）")

    t2 = time.time()
    draft = await run_writing_agent(
        report=report,
        target_word_count=2000,
        writing_architecture=writing_arch,
        strategy_context=strategy,
        project_id=ctx.project_id,
    )
    elapsed2 = time.time() - t2

    # 清理 code fence（agent 已加保護，但雙重保險）
    draft.content_markdown = clean_code_fences(draft.content_markdown)

    # SEO 初檢 → QA 修正 → SEO 複檢（關鍵改動：把失敗項目傳給 QA）
    info("SEO 初檢…")
    pre_seo = run_seo_check_agent(
        draft=draft,
        primary_keyword=primary_kw,
        secondary_keywords=secondary_kwds,
    )
    failed_checks = [c for c in pre_seo["checks"] if not c["passed"]]
    info(f"初檢 {pre_seo['score']} 分（{len(failed_checks)} 項待修：{', '.join(c['name'] for c in failed_checks[:4])}）")

    info("執行 SEO QA 針對性修正…（約 3~8 秒）")
    t2b = time.time()
    draft = await run_seo_qa_agent(
        draft=draft,
        report=report,
        primary_keyword=primary_kw,
        secondary_keywords=secondary_kwds,
        failed_checks=failed_checks,
        project_id=ctx.project_id,
    )
    elapsed2b = time.time() - t2b

    seo_report = run_seo_check_agent(
        draft=draft,
        primary_keyword=primary_kw,
        secondary_keywords=secondary_kwds,
    )

    # 內部連結建議（查詢本專案已發布文章）
    _il_suggestions: list[dict] = []
    try:
        conn_il = sqlite3.connect(DB_PATH)
        _il_rows = conn_il.execute(
            "SELECT title, publish_url, primary_keyword, secondary_keywords "
            "FROM articles WHERE project_id=? AND status='published' "
            "AND publish_url IS NOT NULL AND publish_url != ''",
            (ctx.project_id,),
        ).fetchall()
        conn_il.close()
        _il_existing = [
            {"title": r[0], "url": r[1], "primary_keyword": r[2] or "", "secondary_keywords": r[3] or ""}
            for r in _il_rows if r[1]
        ]
        _il_suggestions = suggest_internal_links(draft.content_markdown, primary_kw, _il_existing)
    except Exception as _e:
        warn(f"內部連結建議失敗：{_e}")

    tick(f"寫文完成！耗時 {elapsed2:.1f}s")
    tick(f"SEO QA 完成！耗時 {elapsed2b:.1f}s")
    tick(f"SEO 檢查：{seo_report['score']} 分（{seo_report['passed_count']}/{seo_report['total_count']}）")
    tick(f"文章標題：{draft.title}")
    tick(f"URL Slug：{draft.slug or '(未生成)'}")
    tick(f"字數：{draft.word_count} 字")
    tick(f"Meta title：{draft.meta_title}")
    tick(f"Meta description：{draft.meta_description[:60]}…")
    if draft.faq_schema_json:
        import json as _json
        _faq_items = _json.loads(draft.faq_schema_json).get("mainEntity", [])
        tick(f"FAQ JSON-LD：{len(_faq_items)} 個問答已產出")
    else:
        info("FAQ JSON-LD：未提取到 FAQ 段落")
    if _il_suggestions:
        tick(f"內部連結建議：{len(_il_suggestions)} 條")
        for _il in _il_suggestions:
            info(f"  「{_il['anchor_text']}」→ {_il['target_title']}")
    else:
        info("內部連結：暫無已發布文章可配對（發布後即可自動配對）")

    # 預覽前 200 字
    preview = draft.content_markdown[:200].replace('\n', ' ')
    info(f"內文預覽：{preview}…")

    # ══════════════════════════════════════════════════════════
    # STEP 3 — 事實查核
    # ══════════════════════════════════════════════════════════
    banner(3, "內容查核 × 法規合規驗證")

    info(f"呼叫 {settings.llm_lite_model} 進行內容合規查核…（約 30 秒）")

    art_type = article.get("article_type", "educational")
    t3 = time.time()
    checked_draft = await run_factcheck_agent(
        draft=draft, report=report, project_id=ctx.project_id, article_type=art_type
    )
    elapsed3 = time.time() - t3

    items = checked_draft.fact_check_items or []
    passed  = [i for i in items if not i.needs_review]
    flagged = [i for i in items if i.needs_review]

    tick(f"查核完成！耗時 {elapsed3:.1f}s")
    tick(f"共查核 {len(items)} 項宣稱")
    tick(f"通過 {GREEN}{len(passed)} 項{RESET}  |  需人工審核 {RED}{len(flagged)} 項{RESET}")

    print()
    for item in items:
        icon = f"{GREEN}🟢{RESET}" if not item.needs_review else f"{RED}🔴{RESET}"
        claim = (item.claim or "")[:55]
        note = (item.reviewer_note or "")[:60] or "OK"
        evidence = ", ".join(item.supporting_evidence)[:50] if item.supporting_evidence else ""
        print(f"     {icon}  {claim}")
        if note and note != "OK":
            print(f"         → {note}")
        if evidence:
            print(f"         証: {evidence}")

    # ══════════════════════════════════════════════════════════
    # STEP 4 — 取得成果
    # ══════════════════════════════════════════════════════════
    banner(4, "存檔 & 輸出")

    # 組裝 factcheck JSON（與 UI 存法一致）
    factcheck_data = {
        "items": [
            {
                "claim": i.claim,
                "reviewer_note": i.reviewer_note,
                "needs_review": i.needs_review,
                "confidence": i.confidence,
                "supporting_evidence": i.supporting_evidence,
            }
            for i in items
        ]
    }

    # 存回資料庫
    save_to_db(
        article_id=article["id"],
        report_json=report.model_dump_json(indent=2),
        draft_content=checked_draft.content_markdown,
        draft_title=draft.title,
        status="reviewing",
        primary_keyword=primary_kw,
        secondary_keywords="\n".join(secondary_kwds),
        slug=draft.slug or "",
        meta_title=draft.meta_title or "",
        meta_description=draft.meta_description or "",
        faq_schema_json=draft.faq_schema_json or "",
        article_schema_json=draft.article_schema_json or "",
        seo_score=seo_report["score"],
    )
    tick("已存回 SQLite 資料庫（研究報告 + 草稿 + 查核結果）")

    # 輸出 Markdown 檔案
    file_slug = re.sub(r'[^\w\u4e00-\u9fff]', '_', primary_kw)[:20] or "article"
    ts   = datetime.now().strftime("%Y%m%d_%H%M")
    md_path = OUTPUT_DIR / f"{ts}_{file_slug}.md"

    # 組裝 FAQ JSON-LD 區塊（貼入 CMS <head> 使用）
    faq_schema_block = ""
    if draft.faq_schema_json:
        faq_schema_block = (
            "\n\n<!-- FAQ Schema JSON-LD（複製此區塊貼入 CMS 的 Schema/Head 欄位）:\n"
            "<script type=\"application/ld+json\">\n"
            f"{draft.faq_schema_json}\n"
            "</script>\n-->"
        )

    # 組裝 Article/BlogPosting JSON-LD 區塊
    article_schema_block = ""
    if draft.article_schema_json:
        article_schema_block = (
            "\n\n<!-- Article Schema JSON-LD（複製此區塊貼入 CMS 的 Schema/Head 欄位）:\n"
            "<script type=\"application/ld+json\">\n"
            f"{draft.article_schema_json}\n"
            "</script>\n-->"
        )

    # 組裝內部連結建議區塊
    internal_links_block = ""
    if _il_suggestions:
        _il_lines = ["\n<!-- 建議內部連結（請在文章內對應錨文字加上超連結）："]
        for _il in _il_suggestions:
            _il_lines.append(f"  錨文字: \"{_il['anchor_text']}\"")
            _il_lines.append(f"  目標: {_il['target_url']}")
            _il_lines.append(f"  標題: {_il['target_title']}")
            _il_lines.append("")
        _il_lines.append("-->")
        internal_links_block = "\n".join(_il_lines)

    md_content = f"""---
title: {draft.title}
slug: {draft.slug or file_slug}
meta_title: {draft.meta_title}
meta_description: {draft.meta_description}
primary_keyword: {primary_kw}
word_count: {draft.word_count}
generated_at: {datetime.now().isoformat()}
factcheck_passed: {len(passed)}/{len(items)}
seo_score: {seo_report['score']}
seo_passed: {seo_report['passed_count']}/{seo_report['total_count']}
---

{checked_draft.content_markdown}{faq_schema_block}{article_schema_block}{internal_links_block}

---
*本文章由 ContentFlow AI 自動生成，建議人工最終審閱後發佈。*
"""
    md_path.write_text(md_content, encoding="utf-8")
    tick(f"Markdown 已輸出 → {md_path.relative_to(ROOT)}")

    # ── 輸出 DOCX ─────────────────────────────────────────────
    docx_path = OUTPUT_DIR / f"{ts}_{file_slug}.docx"
    try:
        export_docx(checked_draft.content_markdown, draft, items, docx_path)
        tick(f"Word 已輸出  → {docx_path.relative_to(ROOT)}")
    except Exception as e:
        warn(f"DOCX 輸出失敗：{e}")
        docx_path = None

    # ── 輸出 PDF ──────────────────────────────────────────────
    pdf_path = OUTPUT_DIR / f"{ts}_{file_slug}.pdf"
    try:
        export_pdf(checked_draft.content_markdown, draft, items, pdf_path)
        tick(f"PDF  已輸出  → {pdf_path.relative_to(ROOT)}")
    except Exception as e:
        warn(f"PDF 輸出失敗（{type(e).__name__}）：{e}")
        warn("如需 PDF 請執行：brew install pango")
        pdf_path = None

    # 統計
    total = time.time() - t0
    print(f"\n{BOLD}{GREEN}{'='*60}{RESET}")
    print(f"{BOLD}{GREEN}  Pipeline 完成！總耗時 {total:.1f}s{RESET}")
    print(f"{GREEN}{'='*60}{RESET}")
    print(f"\n  研究耗時：{elapsed1:.1f}s")
    print(f"  寫文耗時：{elapsed2:.1f}s")
    print(f"  SEO QA：{elapsed2b:.1f}s")
    print(f"  查核耗時：{elapsed3:.1f}s")
    print(f"\n  字數：{draft.word_count} 字")
    print(f"  URL Slug：{draft.slug or '(未生成)'}")
    print(f"  SEO：{seo_report['score']} 分（{seo_report['passed_count']}/{seo_report['total_count']}）")
    print(f"  查核：{len(passed)}/{len(items)} 項通過")
    print(f"\n  SEO 檢查：")
    for check in seo_report['checks']:
        icon = '✓' if check['passed'] else '!'
        print(f"    {icon} {check['detail']}")
    print(f"\n  輸出檔案：")
    print(f"    📄 Markdown → {md_path}")
    if docx_path:
        print(f"    📝 Word     → {docx_path}")
    if pdf_path:
        print(f"    🔖 PDF      → {pdf_path}")
    print()

    return md_path


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="ContentFlow Article Pipeline")
    parser.add_argument("--seqno", type=int, default=4, help="文章序號（預設 4）")
    parser.add_argument("--project", type=str, default="hanben", help="專案 slug（預設 hanben）")
    args = parser.parse_args()
    asyncio.run(run_pipeline(args.seqno, project_slug=args.project))
